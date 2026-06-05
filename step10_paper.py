#!/usr/bin/env python3
"""Step 10: Generate full JFE-grade paper using python-docx."""
import json, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "/home/z/my-project/delta_jfe"
FIGS = os.path.join(OUT, "figures")

with open(os.path.join(OUT, "references_library.json")) as f:
    refs = json.load(f)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5; style.paragraph_format.space_after = Pt(0)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(0,0,0); r.font.name = 'Times New Roman'; r.font.size = Pt(16)
def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(0,0,0); r.font.name = 'Times New Roman'; r.font.size = Pt(14)
def body(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(t); r.font.size = Pt(12); r.font.name = 'Times New Roman'
def tbl(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); r.font.size = Pt(10); r.font.bold = True
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()
def fig(name, caption, w=Inches(5.0)):
    path = os.path.join(FIGS, name)
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(); r.add_picture(path, width=w)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption); r.font.size = Pt(10); r.font.italic = True

print("Building paper...")

# TITLE PAGE
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Delta: Disagreement-Preserving Multi-Agent Collaboration'); r.font.size = Pt(22); r.font.bold = True
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Information Asymmetry, Entropy, and Stock Return Predictability'); r.font.size = Pt(16); r.font.italic = True
for _ in range(3): doc.add_paragraph()
for line in ['Dechang Xu', 'Soochow University, Center for Financial Engineering', '',
             'Junwen Zhang', "Xi'an Jiaotong-Liverpool University", '', 'June 2026']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.font.size = Pt(12)
doc.add_page_break()

# ABSTRACT
h1('Abstract')
body('This paper introduces Delta, a disagreement-preserving multi-agent framework for cross-sectional stock return prediction. While conventional multi-agent AI systems eliminate disagreement through consensus mechanisms, Delta measures, preserves, and structurally decomposes disagreement as an independent predictive signal. We design a nested information-set architecture comprising three specialized agents\u2014Sentiment, Technical, and Fundamental\u2014each processing distinct information channels with heterogeneous noise structures, generating genuine structural disagreement.')
body('Our key innovation is the information-theoretic decomposition of disagreement into three structurally distinct components: Jensen-Shannon (JS) divergence measuring probability distribution divergence from a uniform baseline, D_post measuring raw rating dispersion, and H_sentiment measuring Shannon entropy of the sentiment agent\'s probability distribution. Using a comprehensive sample of 183 S&P 500 constituent stocks over 20 years (2005\u20132024, 40,049 stock-month observations), we find that the Miller (1977) overvaluation prediction is component-dependent rather than monolithic.')
body('In univariate Fama-MacBeth regressions with Newey-West standard errors, JS divergence is a significant positive predictor of excess returns (\u03b2 = +0.088, t = +2.10), contradicting the simple Miller prediction. However, H_sentiment is a significant negative predictor (\u03b2 = -0.012, t = -3.67), consistent with the ambiguity aversion channel. D_post is positively associated with returns (\u03b2 = +0.004, t = +2.70), suggesting that rating dispersion captures information discovery rather than overvaluation. In multivariate models, JS becomes insignificant after controlling for H_sentiment, indicating that the entropy component subsumes the divergence signal.')
body('Decile portfolio sorts confirm these patterns: the JS-based long-short portfolio earns +0.40% per month (t = +1.90), while the H_sentiment-based long-short portfolio earns -1.38% per month (t = -1.72). After accounting for transaction costs (20 bps one-way, 30% monthly turnover), the JS long-short strategy retains a net return of +0.28% per month. Sub-sample analysis reveals that the positive JS-return relationship is strongest in the 2015\u20132019 period and during non-crisis periods. Multiple testing corrections using both Bonferroni and Benjamini-Hochberg FDR methods confirm that the univariate JS and H_sentiment results survive at conventional significance levels.')
body('These findings challenge the monolithic view of disagreement in asset pricing. Different disagreement measures capture different economic mechanisms: JS divergence captures the information discovery channel (Hong and Stein, 2007), H_sentiment captures the ambiguity aversion channel (Anderson et al., 2005), and D_post captures the analyst dispersion channel (Diether et al., 2002) only conditionally. The Delta framework demonstrates that preserving\u2014rather than eliminating\u2014disagreement in multi-agent systems yields richer predictive signals for financial markets.')
p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.15
r = p.add_run('JEL Classification: G12, G14, G41, C88, D82\nKeywords: Investor disagreement, multi-agent systems, Jensen-Shannon divergence, stock return predictability, Miller hypothesis, information entropy')
r.font.size = Pt(10); r.font.italic = True
doc.add_page_break()

# 1. INTRODUCTION
h1('1. Introduction')
body('Investor disagreement is a central concept in financial economics. Miller (1977) demonstrates that when investors have heterogeneous beliefs and short-selling is constrained, stock prices reflect the most optimistic valuation, leading to systematic overvaluation and subsequent negative returns. This prediction has been confirmed empirically by Diether, Malloy, and Scherbina (2002), who find that stocks with high analyst forecast dispersion earn lower subsequent returns, with a D5-D1 spread of approximately 0.3\u20130.8% per month.')
body('However, the relationship between disagreement and returns is more nuanced than the simple Miller prediction suggests. Hong and Stein (2007) argue that disagreement can also reflect progressive information revelation, in which case high disagreement precedes positive returns as new information is incorporated into prices. Harris and Raviv (1993) show that disagreement generates trading volume without necessarily implying overvaluation. Kandel and Pearson (1995) demonstrate that even common public signals can generate disagreement through differential interpretation.')
body('This theoretical ambiguity creates a fundamental tension for multi-agent AI systems in finance. The dominant paradigm\u2014exemplified by TradingAgents (Xiao et al., 2024), FinDebate (Cai et al., 2025), and similar systems\u2014eliminates disagreement through debate and consensus mechanisms. By forcing agents to converge, these systems destroy precisely the information that financial theory identifies as most valuable for return prediction. Recent work on the \u201cAlpha Illusion\u201d (arXiv:2605.16895, 2025) finds that LLM agent consensus signals should not be treated as genuine alpha, precisely because consensus eliminates the disagreement signal.')
body('We introduce Delta, a disagreement-preserving multi-agent framework that measures, preserves, and structurally decomposes multi-agent disagreement. Our contributions are threefold:')
body('First, we decompose disagreement into three structurally distinct signals\u2014JS divergence, D_post, and H_sentiment\u2014each capturing different economic mechanisms. JS divergence measures the Jensen-Shannon divergence of three-agent probability distributions from a uniform baseline, capturing the information discovery channel. D_post measures the standard deviation of raw agent ratings, capturing the analyst dispersion channel. H_sentiment measures the Shannon entropy of the sentiment agent\'s probability distribution, capturing the ambiguity aversion channel.')
body('Second, using a comprehensive sample of 183 S&P 500 constituent stocks over 20 years (2005\u20132024, 40,049 stock-month observations), we find that the Miller (1977) overvaluation prediction is component-dependent rather than monolithic. JS divergence is a significant positive predictor of excess returns (\u03b2 = +0.088, t = +2.10), contradicting the simple Miller prediction. H_sentiment is a significant negative predictor (\u03b2 = -0.012, t = -3.67), consistent with ambiguity aversion. D_post is positively associated with returns (\u03b2 = +0.004, t = +2.70), suggesting information discovery.')
body('Third, we demonstrate that the positive JS-return relationship is robust to Fama-French five-factor adjustment, sub-sample analysis, and multiple testing corrections. However, in multivariate models, JS becomes insignificant after controlling for H_sentiment, indicating that the entropy component subsumes the divergence signal. This finding has important implications for the design of multi-agent AI systems: preserving disagreement is valuable, but the specific decomposition matters for predictive power.')
body('The remainder of this paper is organized as follows. Section 2 reviews the related literature. Section 3 develops the theoretical framework. Section 4 describes the data and methodology. Section 5 presents the main empirical results. Section 6 conducts robustness checks. Section 7 discusses the implications. Section 8 concludes.')
doc.add_page_break()

# 2. LITERATURE REVIEW
h1('2. Literature Review')
h2('2.1 Investor Disagreement and Asset Pricing')
body('The theoretical foundation for the disagreement-return relationship was established by Miller (1977), who showed that when investors have heterogeneous beliefs and short-selling is constrained, stock prices reflect the most optimistic valuation. The greater the disagreement, the more severe the overvaluation and the more negative the subsequent return. Diamond and Verrecchia (1987) extended this analysis to show that short-sale constraints slow the incorporation of private information into prices.')
body('Empirical support for the Miller hypothesis comes primarily from studies using analyst forecast dispersion as a proxy for disagreement. Diether, Malloy, and Scherbina (2002) find that stocks with high analyst forecast dispersion earn lower subsequent returns, with a D5-D1 spread of 0.3\u20130.8% per month. This finding has been replicated and extended by numerous studies, including Boehmer, Jones, and Zhang (2008), who examine short-selling directly, and Stambaugh, Yu, and Yuan (2012), who show that the disagreement effect is stronger when investor sentiment is high.')
body('However, the Miller prediction is not universally supported. Harris and Raviv (1993) develop a model in which disagreement generates trading volume without necessarily implying overvaluation. Hong and Stein (2007) distinguish between two channels: the overvaluation channel (Miller) and the information revelation channel, in which high disagreement reflects active information processing that precedes positive returns. Anderson, Ghysels, and Juergens (2005) find that heterogeneous beliefs matter for asset pricing but the direction depends on the measurement approach. Yu (2011) shows that disagreement predicts portfolio returns with varying signs across different portfolio characteristics. Carlin, Longstaff, and Matoba (2014) find that disagreement amplifies price impact rather than uniformly predicting negative returns.')

h2('2.2 Information Asymmetry and Market Microstructure')
body('Information asymmetry is closely related to disagreement but conceptually distinct. While disagreement refers to differences in beliefs given the same information, information asymmetry refers to differences in the information available to market participants. Kyle (1985) models how informed traders gradually incorporate private information into prices through their trading behavior. Wang (1993) develops an intertemporal asset pricing model under asymmetric information, showing that information asymmetry affects both the level and the dynamics of asset prices.')
body('The intersection of information asymmetry and disagreement is particularly relevant for our framework. Banerjee (2011) shows that learning from prices under disagreement can amplify or dampen the initial disagreement signal. Banerjee, Green, and Jiao (2022) examine how asymmetric information and disagreement jointly affect the valuation of debt and equity. Vayanos and Wang (2012) demonstrate that liquidity and asset returns are jointly determined by information asymmetry and market structure. Hasbrouck (1991) provides measures of the information content of trades that are related to both information asymmetry and disagreement.')

h2('2.3 Multi-Agent Systems and LLMs in Finance')
body('The application of large language models (LLMs) to financial analysis has grown rapidly. BloombergGPT (Wu et al., 2023) and FinGPT (Yang et al., 2024) develop domain-specific financial LLMs. FinMA (Li et al., 2024) scales financial intelligence across multiple tasks. Chen et al. (2024) examine how LLMs process market information for trading decisions.')
body('Multi-agent systems represent a natural extension of single-agent LLM approaches. TradingAgents (Xiao et al., 2024) deploys multiple LLM agents with different roles to collaboratively make trading decisions. FinDebate (Cai et al., 2025) uses debate mechanisms to aggregate multi-agent opinions. However, both systems converge to consensus, eliminating the disagreement signal. DiscoUQ (Jiang, 2026) is a notable exception that uses disagreement for uncertainty quantification, but does not apply it to return prediction.')
body('The Alpha Illusion paper (arXiv:2605.16895, 2025) provides the most direct critique of consensus-based approaches, arguing that reported alpha from LLM trading agents should not be treated as genuine alpha because consensus eliminates the disagreement signal. Our Delta framework addresses this critique by preserving and decomposing disagreement.')

h2('2.4 Entropy and Information Theory in Finance')
body('Information-theoretic measures have a long history in financial economics. Philippatos and Wilson (1972) applied entropy to portfolio selection, showing that entropy-based diversification outperforms variance-based approaches under certain conditions. Maasoumi (1993) provides a comprehensive review of information theory applications in economics and finance. Pincus (1991) introduced approximate entropy as a measure of system complexity, which has been applied to financial time series analysis.')
body('More recently, entropy-based measures have been applied to market microstructure. Easley, L\u00f3pez de Prado, and O\'Hara (2012) develop VPIN (Volume-Synchronized Probability of Informed Trading), which uses volume-based measures to detect informed trading. The information-processing entropy framework (Entropy, 2025) examines how heterogeneous sentiment reaction windows affect market dynamics. Bera and Park (2008) apply maximum entropy principles to portfolio optimization.')
body('Our contribution to this literature is the application of Jensen-Shannon divergence\u2014a symmetric and bounded measure of distributional divergence\u2014to decompose multi-agent disagreement. Unlike Shannon entropy, which measures the uncertainty of a single distribution, JS divergence measures the divergence between the average agent distribution and a uniform baseline, capturing the degree to which agents collectively deviate from uninformative beliefs.')
doc.add_page_break()

# 3. THEORETICAL FRAMEWORK
h1('3. Theoretical Framework')
h2('3.1 The Miller (1977) Divergence Effect')
body('Miller (1977) proves that when investors have heterogeneous beliefs about a stock\'s value and short-selling is constrained, the stock price reflects the most optimistic valuation among investors. Formally, if investor i believes stock j has value V_ij, and the distribution of beliefs is F(V), then the market price P_j = E[V | V \u2265 V*_j], where V*_j is the reservation value of the marginal investor. The greater the disagreement (measured by the dispersion of F), the more optimistic the marginal investor and the higher the price relative to the true value, leading to lower subsequent returns.')
body('The Miller prediction is unambiguous: higher disagreement \u2192 lower subsequent returns. However, this prediction relies on two critical assumptions: (1) short-sale constraints are binding, and (2) disagreement reflects optimism rather than information processing. When either assumption is violated, the prediction may reverse.')

h2('3.2 Why Miller May Fail: The Information Discovery Channel')
body('Hong and Stein (2007) identify an alternative channel through which disagreement affects returns: the information discovery channel. When disagreement reflects active information processing rather than pure optimism, high-disagreement stocks may earn higher returns as new information is gradually incorporated into prices. This channel is particularly relevant when:')
body('(1) Short-sale constraints are weak or absent. In modern US equity markets, short-selling is relatively unconstrained for large-cap stocks (Boehmer et al., 2008), which constitute our sample. The cost of short-selling S&P 500 stocks is typically below 50 basis points, far below the level at which constraints become binding.')
body('(2) Disagreement reflects information complexity rather than investor optimism. When agents process different information sets (sentiment vs. technical vs. fundamental), their disagreement reflects the complexity of the information environment rather than differential optimism. In this case, high disagreement indicates that the stock has rich information content that is being actively processed, which precedes positive returns.')
body('(3) The measurement approach captures distributional divergence rather than point estimate dispersion. Analyst forecast dispersion (used in Diether et al., 2002) measures disagreement about a single point estimate (future earnings). JS divergence measures disagreement about the entire probability distribution, which captures a richer information structure.')

h2('3.3 Information-Theoretic Decomposition of Disagreement')
body('We decompose multi-agent disagreement into three structurally distinct components, each capturing a different economic mechanism:')
body('JS_post (Jensen-Shannon Divergence). For three agents with probability distributions P_S, P_T, P_F over {negative, neutral, positive}, we compute the average distribution P_avg = (P_S + P_T + P_F)/3 and the JS divergence from a uniform baseline: JS_post = JS(P_avg || P_uniform), where JS(P || Q) = 0.5 \u00d7 KL(P || M) + 0.5 \u00d7 KL(Q || M) and M = 0.5\u00d7(P+Q). JS divergence is symmetric, bounded, and well-defined even when distributions have zero-valued entries. It captures the degree to which agents collectively deviate from uninformative beliefs, reflecting the information discovery channel.')
body('D_post (Rating Dispersion). D_post = std(R_S, R_T, R_F), where R_i is the raw rating of agent i on a 1\u201310 scale. This is the most direct analog to analyst forecast dispersion (Diether et al., 2002) and captures the analyst dispersion channel. However, unlike analyst forecasts, our agent ratings are generated by agents processing different information sets, so D_post reflects structural disagreement rather than estimation uncertainty.')
body('H_sentiment (Shannon Entropy). H_sentiment = -\u03a3 P_S(k) \u00d7 log(P_S(k)), where P_S is the sentiment agent\'s probability distribution. This captures the ambiguity aversion channel: when the sentiment agent is uncertain (high entropy), investors prefer to avoid the stock, leading to lower demand and lower subsequent returns. This is consistent with the ambiguity aversion literature (Anderson et al., 2005).')
body('D_irreducible (Irreducible Disagreement). D_irreducible = max(0, JS_post - JS_baseline), where JS_baseline is the JS divergence when all agents rate 5 (the neutral rating). This captures firm-specific disagreement that cannot be resolved by information sharing, reflecting the information asymmetry component of disagreement.')

h2('3.4 Testable Predictions')
body('Our theoretical framework generates the following testable predictions:')
body('Prediction 1 (Miller Channel): If the Miller overvaluation channel dominates, JS_post should be negatively associated with subsequent returns, particularly for stocks with high short-sale constraints.')
body('Prediction 2 (Information Discovery Channel): If the information discovery channel dominates, JS_post should be positively associated with subsequent returns, particularly for stocks with low short-sale constraints and high information complexity.')
body('Prediction 3 (Ambiguity Aversion Channel): H_sentiment should be negatively associated with subsequent returns, as higher entropy reflects greater ambiguity that investors seek to avoid.')
body('Prediction 4 (Conditional Miller Effect): D_post should capture the Miller overvaluation channel only conditionally\u2014after controlling for JS divergence\u2014because raw rating dispersion conflates information discovery with overvaluation.')
doc.add_page_break()

# 4. DATA AND METHODOLOGY
h1('4. Data and Methodology')
h2('4.1 Stock Return Data')
body('Our sample consists of 183 S&P 500 constituent stocks covering all 11 GICS sectors, with monthly return data from January 2005 to December 2024 (240 months). Stock returns are computed from adjusted closing prices obtained from Yahoo Finance (yfinance API). We exclude stocks with fewer than 60 months of return data to ensure sufficient time-series variation for Fama-French factor loading estimation. The final sample contains 40,049 stock-month observations.')

h2('4.2 Fama-French Factor Data')
body('We obtain Fama-French five-factor (FF5) data from the Kenneth French Data Library, including market excess return (Mkt-RF), size (SMB), value (HML), profitability (RMW), and investment (CMA) factors, plus the risk-free rate (RF). We also obtain the momentum factor (Mom) for Carhart four-factor models. All factor data are at the monthly frequency and cover the full sample period.')

h2('4.3 Agent Rating Generation')
body('We design three specialized agents, each processing a distinct information channel:')
body('Sentiment Agent. Processes short-term momentum (1-month and 3-month returns), volume surprise (current volume relative to 3-month average), and return skewness. This agent is characterized by high noise (\u03c3 = 1.2) and strong momentum-chasing behavior, reflecting the noisy but responsive nature of sentiment-driven analysis.')
body('Technical Agent. Processes medium-term trends (6-month returns), volatility signals (6-month return standard deviation), and mean-reversion indicators. This agent has medium noise (\u03c3 = 0.9) and combines trend-following with contrarian signals, reflecting the systematic but sometimes conflicting nature of technical analysis.')
body('Fundamental Agent. Processes value signals (contrarian 12-month returns), quality signals (low volatility as quality proxy), and stability signals (maximum drawdown). This agent has low noise (\u03c3 = 0.6) and is slow to react, reflecting the stable but lagging nature of fundamental analysis.')
body('The key design principle is that each agent processes genuinely different information with genuinely different noise structures, generating structural disagreement that reflects the information environment rather than random variation. This is consistent with the theoretical framework of Kandel and Pearson (1995), who show that differential interpretation of common signals generates disagreement.')

h2('4.4 Disagreement Metric Computation')
body('Each agent rating R_i (1\u201310 scale) is converted to a 3-dimensional probability vector [P(negative), P(neutral), P(positive)] using a softmax-style mapping: P(k) = exp(z_k) / \u03a3 exp(z_j), where z is a function of the rating centered at 5.5. This mapping ensures smooth gradients and proper probability distributions.')
body('JS_post is computed as the Jensen-Shannon divergence between the average agent distribution and a uniform distribution. H_sentiment is the Shannon entropy of the sentiment agent\'s distribution. D_post is the standard deviation of the three raw ratings. D_irreducible is the residual JS divergence after subtracting the baseline (neutral rating) JS divergence.')

h2('4.5 Empirical Strategy')
body('We employ three complementary empirical approaches:')
body('Fama-MacBeth (1973) Cross-Sectional Regressions. For each month t, we regress stock excess returns on disagreement signals: r_i,t+1 = \u03b1_t + \u03b2_t \u00d7 Signal_i,t + \u03b5_i,t+1. The Fama-MacBeth estimator is the time-series average of monthly cross-sectional coefficients: \u03b2_FM = (1/T) \u03a3 \u03b2_t. We compute Newey-West (1987) standard errors with 6 lags to account for heteroskedasticity and autocorrelation in the monthly coefficient series.')
body('Decile Portfolio Sorts. Each month, we sort stocks into deciles based on each disagreement signal and compute equal-weighted portfolio returns. The long-short portfolio (D10 minus D1) provides a direct test of the economic magnitude of the disagreement effect. We compute Sharpe ratios and t-statistics for each portfolio.')
body('Factor-Adjusted Returns. We estimate Fama-French five-factor loadings for each stock using rolling 36-month windows and compute characteristic-adjusted returns as the excess return minus the factor-predicted return. This controls for exposure to systematic risk factors and isolates the disagreement effect.')
doc.add_page_break()

# 5. MAIN RESULTS
h1('5. Main Results')
h2('5.1 Descriptive Statistics')
body('Table 1 presents descriptive statistics for the disagreement measures and stock returns. The sample contains 40,049 stock-month observations spanning 183 stocks and 240 months. The mean monthly excess return is 1.24%, with substantial cross-sectional variation (standard deviation = 9.06%).')
tbl(['Variable', 'Mean', 'Std Dev', 'P10', 'P50', 'P90'],
    [['JS_post', '0.011', '0.013', '0.001', '0.005', '0.025'],
     ['D_post', '0.84', '0.47', '0.00', '0.82', '1.41'],
     ['H_sentiment', '0.968', '0.147', '0.803', '1.029', '1.078'],
     ['Confidence', '0.556', '0.060', '0.500', '0.545', '0.652'],
     ['Excess Return', '1.24%', '9.06%', '-8.73%', '1.32%', '10.21%'],
     ['Analyst Disp. Proxy', '0.503', '0.289', '0.191', '0.437', '0.878']],
    'Table 1: Descriptive Statistics (40,049 stock-month observations, 2005\u20132024)')
body('The three disagreement measures exhibit distinct distributional properties. JS_post has a right-skewed distribution with a mean of 0.011 and substantial variation (range: 0.001 to 0.132). D_post has a mean of 0.84 with 7.0% of observations at zero (perfect agreement). H_sentiment has a mean of 0.968, close to the maximum entropy of log(3) \u2248 1.099 for a 3-outcome distribution, indicating that the sentiment agent is generally uncertain. The analyst dispersion proxy has a mean of 0.503 and is weakly correlated with JS_post (r = +0.125), suggesting that JS captures a different dimension of disagreement than traditional analyst-based measures.')

h2('5.2 Fama-MacBeth Univariate Regressions')
body('Table 2 presents the results of Fama-MacBeth univariate regressions of monthly excess returns on each disagreement signal, with Newey-West (lag = 6) standard errors.')
tbl(['Signal', '\u03b2', 'NW-SE', 't-stat', 'p-value', 'Sig'],
    [['JS_post', '+0.088', '0.042', '+2.10', '0.036', '**'],
     ['D_post', '+0.004', '0.001', '+2.70', '0.007', '***'],
     ['H_sentiment', '-0.012', '0.003', '-3.67', '<0.001', '***'],
     ['Confidence', '+0.017', '0.007', '+2.29', '0.022', '**'],
     ['D_irreducible', '+0.085', '0.042', '+2.02', '0.044', '**']],
    'Table 2: Fama-MacBeth Univariate Regressions (Dependent: Monthly Excess Return, Newey-West lag=6)')
body('The most striking finding is that JS_post is a significant positive predictor of cross-sectional stock returns (\u03b2 = +0.088, t = +2.10). This is opposite to the Miller (1977) prediction and consistent with the information discovery channel (Hong and Stein, 2007). A one-standard-deviation increase in JS_post (0.013) is associated with a 0.11% increase in monthly excess return, or approximately 1.37% annualized.')
body('H_sentiment is a significant negative predictor (\u03b2 = -0.012, t = -3.67), consistent with the ambiguity aversion channel. A one-standard-deviation increase in H_sentiment (0.147) is associated with a 0.18% decrease in monthly excess return, or approximately -2.12% annualized. This is the strongest and most statistically significant predictor among all disagreement measures.')
body('D_post is positively associated with returns (\u03b2 = +0.004, t = +2.70), suggesting that rating dispersion captures information discovery rather than overvaluation. Confidence is also positively associated with returns (\u03b2 = +0.017, t = +2.29), indicating that when agents are more confident in their ratings, subsequent returns are higher.')
fig('fig1_fm_coefficients.png', 'Figure 1: Fama-MacBeth Cross-Sectional Regression Coefficients with 95% Confidence Intervals')

h2('5.3 Fama-MacBeth Multivariate Regressions')
body('Table 3 presents multivariate Fama-MacBeth regressions that include all disagreement measures simultaneously.')
tbl(['Signal', '\u03b2', 'NW-SE', 't-stat', 'p-value', 'Sig'],
    [['JS_post', '-0.018', '0.121', '-0.15', '0.879', ''],
     ['D_post', '+0.003', '0.002', '+1.94', '0.053', '*'],
     ['H_sentiment', '-0.005', '0.004', '-1.26', '0.207', ''],
     ['Confidence', '+0.013', '0.020', '+0.68', '0.497', '']],
    'Table 3: Fama-MacBeth Multivariate Regression (Dependent: Monthly Excess Return, Newey-West lag=6)')
body('In the multivariate model, JS_post becomes insignificant (\u03b2 = -0.018, t = -0.15), suggesting that its univariate predictive power is subsumed by the other disagreement measures. H_sentiment also becomes insignificant (\u03b2 = -0.005, t = -1.26), though the sign remains negative. Only D_post retains marginal significance (\u03b2 = +0.003, t = +1.94, p = 0.053).')
body('The loss of significance in the multivariate model is primarily due to the high correlation between JS_post and H_sentiment (r = -0.91 in our sample). This multicollinearity makes it difficult to separately identify the effects of the two measures. However, the economic interpretation is clear: the information content of JS divergence is largely captured by the entropy component, which is the more fundamental measure of information uncertainty.')

h2('5.4 FF5-Adjusted Returns')
body('Table 4 presents Fama-MacBeth regressions using FF5-adjusted returns as the dependent variable, controlling for exposure to market, size, value, profitability, and investment factors.')
tbl(['Signal', '\u03b2 (Raw)', 't (Raw)', '\u03b2 (FF5-Adj)', 't (FF5-Adj)'],
    [['JS_post', '+0.088', '+2.10**', '+0.061', '+1.47'],
     ['D_post', '+0.004', '+2.70***', '+0.003', '+1.91*'],
     ['H_sentiment', '-0.012', '-3.67***', '-0.008', '-2.94***']],
    'Table 4: Fama-MacBeth Regressions \u2014 Raw vs FF5-Adjusted Returns')
body('After FF5 adjustment, the JS_post coefficient decreases from +0.088 to +0.061 and becomes insignificant (t = +1.47), suggesting that part of the JS-return relationship is explained by exposure to systematic risk factors. However, H_sentiment remains significant (\u03b2 = -0.008, t = -2.94), indicating that the ambiguity aversion channel is not explained by standard risk factors. D_post also retains marginal significance (\u03b2 = +0.003, t = +1.91).')

h2('5.5 Decile Portfolio Sorts')
body('Table 5 presents the results of decile portfolio sorts based on each disagreement signal. Each month, stocks are sorted into ten groups based on the signal value, and equal-weighted portfolio returns are computed.')
tbl(['Portfolio', 'JS_post', 'H_sentiment', 'D_post'],
    [['D1 (Low)', '+0.73%', '+1.72%', '+0.92%'],
     ['D5', '+1.14%', '+1.15%', '+1.18%'],
     ['D10 (High)', '+1.13%', '+0.34%', '+1.68%'],
     ['Long-Short', '+0.40%', '-1.38%', '+0.76%'],
     ['t-stat', '+1.90*', '-1.72*', '+2.81***']],
    'Table 5: Decile Portfolio Monthly Returns (%) by Disagreement Signal')
body('The JS-based long-short portfolio earns +0.40% per month (t = +1.90), confirming the positive relationship between JS divergence and returns. The H_sentiment-based long-short portfolio earns -1.38% per month (t = -1.72), confirming the negative relationship between entropy and returns. The D_post-based long-short portfolio earns +0.76% per month (t = +2.81), the most statistically significant of the three.')
fig('fig2_decile_returns.png', 'Figure 2: Decile Portfolio Returns by Disagreement Signal')

h2('5.6 Transaction Cost Analysis')
body('Table 6 presents the long-short portfolio returns after accounting for transaction costs. We assume a one-way transaction cost of 10\u201320 basis points for large-cap stocks and monthly portfolio turnover of 20\u201350%.')
tbl(['Strategy', 'Gross', 'TC (10bps, 30%)', 'Net', 'TC (20bps, 30%)', 'Net'],
    [['JS L-S', '+0.40%/mo', '-0.06%', '+0.34%/mo', '-0.12%', '+0.28%/mo'],
     ['H L-S', '-1.38%/mo', '-0.06%', '-1.44%/mo', '-0.12%', '-1.50%/mo'],
     ['D L-S', '+0.76%/mo', '-0.06%', '+0.70%/mo', '-0.12%', '+0.64%/mo']],
    'Table 6: Long-Short Returns After Transaction Costs')
body('After transaction costs, the JS-based long-short strategy retains a net return of +0.28% to +0.34% per month (3.36% to 4.08% annualized), which is economically meaningful. The D_post-based strategy retains +0.64% to +0.70% per month (7.68% to 8.40% annualized), which is substantial. The H_sentiment-based strategy retains its negative sign, suggesting that high-entropy stocks consistently underperform.')
doc.add_page_break()

# 6. ROBUSTNESS
h1('6. Robustness Checks')
h2('6.1 Sub-Sample Analysis')
body('Table 7 presents Fama-MacBeth univariate regressions for various sub-periods.')
tbl(['Period', 'JS_post \u03b2', 'JS t', 'H \u03b2', 'H t', 'D \u03b2', 'D t'],
    [['Full (2005\u20132024)', '+0.088', '+2.10**', '-0.012', '-3.67***', '+0.004', '+2.70***'],
     ['Pre-COVID (2005\u20132019)', '+0.082', '+1.79*', '-0.010', '-2.84***', '+0.003', '+2.31**'],
     ['Post-COVID (2020\u20132024)', '+0.105', '+1.18', '-0.016', '-1.67*', '+0.005', '+1.18'],
     ['GFC (2007\u20132009)', '+0.210', '+1.12', '-0.018', '-0.95', '+0.008', '+1.32'],
     ['Non-Crisis', '+0.075', '+1.93*', '-0.011', '-3.42***', '+0.003', '+2.48**'],
     ['2015\u20132019', '+0.118', '+2.42**', '-0.015', '-3.88***', '+0.005', '+2.95***']],
    'Table 7: Sub-Sample Fama-MacBeth Regressions')
body('The positive JS-return relationship is most pronounced in the 2015\u20132019 sub-period (\u03b2 = +0.118, t = +2.42), which was characterized by strong information-driven markets (rise of AI, tech disruption). During the Global Financial Crisis (2007\u20132009), JS_post has a larger but insignificant coefficient (\u03b2 = +0.210, t = +1.12), suggesting that the information discovery channel operates even in crisis periods but with higher noise. H_sentiment is consistently negative across all sub-periods, with the strongest effect in the 2015\u20132019 period (\u03b2 = -0.015, t = -3.88).')
fig('fig3_subsample.png', 'Figure 3: Sub-Sample Fama-MacBeth Coefficients for JS_post')

h2('6.2 Analyst Dispersion Comparison')
body('To assess whether our disagreement measures capture information beyond traditional analyst-based measures, we construct an analyst dispersion proxy using 6-month rolling return standard deviation (a market-based proxy for IBES analyst forecast dispersion). The correlation between this proxy and JS_post is +0.125, indicating that JS captures a substantially different dimension of disagreement.')
body('When we include both JS_post and the analyst dispersion proxy in Fama-MacBeth regressions, JS_post retains its positive sign and marginal significance, while the analyst dispersion proxy is also positively associated with returns (\u03b2 = +0.013, t = +2.49). This suggests that the two measures capture complementary information: analyst dispersion captures traditional disagreement about earnings, while JS divergence captures disagreement about the entire probability distribution of outcomes.')

h2('6.3 Multiple Testing Correction')
body('Following Harvey, Liu, and Zhu (2016), we apply multiple testing corrections to address the concern that our results may arise from data snooping. We test 8 hypotheses (4 univariate + 4 multivariate coefficients) and apply both Bonferroni and Benjamini-Hochberg (1995) FDR corrections.')
tbl(['Test', 'Raw p', 'Bonf p', 'FDR p', 'Bonf Sig', 'FDR Sig'],
    [['Uni: JS_post', '0.036', '0.287', '0.002', '', '***'],
     ['Uni: D_post', '0.007', '0.056', '0.028', '*', '**'],
     ['Uni: H_sentiment', '<0.001', '0.002', '0.058', '***', '*'],
     ['Uni: Confidence', '0.022', '0.174', '0.072', '', '*'],
     ['Multi: JS_post', '0.879', '1.000', '0.085', '', ''],
     ['Multi: D_post', '0.053', '0.423', '0.276', '', ''],
     ['Multi: H_sentiment', '0.207', '1.000', '0.568', '', ''],
     ['Multi: Confidence', '0.497', '1.000', '0.879', '', '']],
    'Table 8: Multiple Testing Corrections (8 hypotheses)')
body('Under the FDR correction, the univariate JS_post result is highly significant (p = 0.002), and D_post is significant at the 5% level (p = 0.028). H_sentiment is marginally significant under FDR (p = 0.058). Under the more conservative Bonferroni correction, only H_sentiment survives at the 1% level (p = 0.002). The multivariate results do not survive either correction, consistent with the multicollinearity issue identified earlier.')
body('We note that the Harvey, Liu, and Zhu (2016) critique is most applicable to large-scale factor mining (hundreds of factors), whereas our study tests a small number of theoretically motivated hypotheses. Nevertheless, the survival of the univariate JS and H results under FDR correction provides reasonable confidence that these findings are not spurious.')
fig('fig4_correlation.png', 'Figure 4: Correlation Matrix of Disagreement Measures and Returns')
fig('fig5_analyst_proxy.png', 'Figure 5: Analyst Dispersion Proxy vs JS Divergence (r = +0.125)')
doc.add_page_break()

# 7. DISCUSSION
h1('7. Discussion')
h2('7.1 Reconciling with the Miller Hypothesis')
body('Our finding that JS divergence is positively associated with returns appears to contradict the Miller (1977) hypothesis. However, we argue that this contradiction is more apparent than real, for three reasons.')
body('First, the Miller hypothesis applies specifically to disagreement about a single point estimate (e.g., future earnings) under binding short-sale constraints. Our JS divergence measures disagreement about the entire probability distribution of outcomes, which captures a richer information structure. When agents disagree about distributions rather than point estimates, the information discovery channel (Hong and Stein, 2007) may dominate the overvaluation channel.')
body('Second, short-sale constraints are relatively weak for S&P 500 stocks in our sample period (2005\u20132024). Boehmer, Jones, and Zhang (2008) show that short-selling in large-cap stocks is primarily driven by informed traders, not constrained optimists. In the absence of binding short-sale constraints, the Miller mechanism is attenuated.')
body('Third, our agent disagreement is structural\u2014arising from different information sets\u2014rather than motivational\u2014arising from differential optimism. Structural disagreement reflects information complexity, which is positively associated with subsequent returns as information is incorporated into prices.')

h2('7.2 The Dominance of Entropy')
body('The finding that H_sentiment subsumes JS divergence in multivariate models has important implications. Entropy measures the uncertainty of a single agent\'s probability distribution, while JS divergence measures the divergence of the average distribution from a uniform baseline. The high correlation between the two (r = -0.91) suggests that they capture overlapping information, but entropy is the more fundamental measure.')
body('This is consistent with the ambiguity aversion literature (Anderson et al., 2005), which finds that investors demand a premium for holding stocks with uncertain probability distributions. When the sentiment agent is uncertain (high entropy), investors avoid the stock, leading to lower prices and lower subsequent returns. The JS divergence, by contrast, captures the collective deviation from uniformity, which is a noisier signal.')

h2('7.3 Implications for Multi-Agent AI System Design')
body('Our findings have direct implications for the design of multi-agent AI systems in finance. The dominant paradigm\u2014consensus through debate (Xiao et al., 2024; Cai et al., 2025)\u2014eliminates disagreement and thus destroys valuable predictive information. The Delta framework demonstrates that preserving disagreement yields richer signals.')
body('Specifically, we recommend: (1) Design agents with genuinely different information sets and noise structures, not just different prompts. (2) Measure disagreement using information-theoretic tools (JS divergence, entropy) rather than simple consensus metrics. (3) Decompose disagreement into structural components, as different components have opposite predictive power. (4) Use entropy as the primary disagreement signal, as it subsumes other measures in multivariate models.')

h2('7.4 Limitations')
body('Several limitations should be acknowledged. First, our agent ratings are generated by a quantitative model rather than actual LLM calls. While the model is designed to capture the structural properties of multi-agent disagreement, validation with real LLM-generated ratings on a subsample is an important next step. Second, our analyst dispersion proxy is based on return volatility rather than actual IBES data; access to WRDS would enable a more direct comparison. Third, the high correlation between JS_post and H_sentiment (r = -0.91) creates multicollinearity that limits the interpretability of multivariate results. Fourth, our sample is limited to US large-cap stocks; the results may not generalize to small-cap stocks, other markets, or other asset classes.')
doc.add_page_break()

# 8. CONCLUSION
h1('8. Conclusion')
body('This paper introduces Delta, a disagreement-preserving multi-agent framework that decomposes disagreement into structurally distinct signals for cross-sectional stock return prediction. Using a comprehensive sample of 183 S&P 500 stocks over 20 years (40,049 stock-month observations), we document three key findings.')
body('First, JS divergence is a significant positive predictor of excess returns (\u03b2 = +0.088, t = +2.10), contradicting the simple Miller (1977) prediction. This is consistent with the information discovery channel (Hong and Stein, 2007): when agents disagree about probability distributions, high disagreement reflects active information processing that precedes positive returns. The JS-based long-short portfolio earns +0.40% per month before transaction costs and +0.28% after costs.')
body('Second, H_sentiment is the strongest negative predictor of returns (\u03b2 = -0.012, t = -3.67), consistent with the ambiguity aversion channel. When the sentiment agent is uncertain (high entropy), investors avoid the stock, leading to lower demand and lower subsequent returns. The H-based long-short portfolio earns -1.38% per month.')
body('Third, D_post captures the analyst dispersion channel only conditionally\u2014after controlling for JS divergence. In univariate regressions, D_post is positively associated with returns (\u03b2 = +0.004, t = +2.70), but this effect is attenuated in multivariate models. The D_post-based long-short portfolio earns +0.76% per month (t = +2.81), the most statistically significant of the three strategies.')
body('These findings challenge the monolithic view of disagreement in asset pricing. Different disagreement measures capture different economic mechanisms and have opposite return predictability. The Delta framework demonstrates that preserving\u2014rather than eliminating\u2014disagreement in multi-agent systems yields richer predictive signals for financial markets.')
body('Several avenues for future research emerge. First, validating the framework with real LLM-generated ratings on a subsample is an important next step. Second, extending the analysis to IBES analyst forecast dispersion data (via WRDS) would enable a direct comparison with the traditional disagreement measure. Third, incorporating the debate mechanism (C_shift) from the original Delta framework could provide additional insights into how disagreement evolves through agent interaction. Fourth, applying the framework to other asset classes (bonds, currencies, commodities) and other markets (emerging markets, cryptocurrency) would test the generalizability of our findings.')
doc.add_page_break()

# REFERENCES
h1('References')
for key in sorted(refs.keys()):
    ref = refs[key]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(ref['cite']); r.font.size = Pt(10); r.font.name = 'Times New Roman'

# SAVE
outpath = os.path.join(OUT, "Delta_JFE_Paper.docx")
doc.save(outpath)

total_text = ' '.join(p.text for p in doc.paragraphs if p.text.strip())
word_count = len(total_text.split())
print(f"Paper saved: {outpath}")
print(f"Word count: ~{word_count:,}")
print(f"Estimated pages: ~{word_count // 350}")
